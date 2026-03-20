from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import matplotlib.pyplot as plt
import base64
from typing import Dict

def markdown_to_docx(content: str, title: str = "") -> bytes:
    doc = Document()
    if title:
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.strip():
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generate_radar_chart(scores: Dict[str, float], title: str = "能力雷达图") -> str:
    categories = list(scores.keys())
    values = list(scores.values())
    values += values[:1]
    
    angles = [n / float(len(categories)) * 2 * 3.14159 for n in range(len(categories))]
    angles += angles[:1]
    
    fig, ax = plt.subplots(subplot_kw=dict(polar=True))
    ax.plot(angles, values, 'o-', linewidth=2)
    ax.fill(angles, values, alpha=0.25)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories)
    ax.set_title(title)
    
    buffer = BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
    buffer.seek(0)
    plt.close()
    return base64.b64encode(buffer.read()).decode()
